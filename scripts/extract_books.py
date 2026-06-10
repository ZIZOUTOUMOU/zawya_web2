"""
Extract book catalog data from .docx library management files in lirbrery_books.
Each .docx file = one Dewey Decimal category with a table of book entries.
Outputs: JSON catalog + SQL seed script for the Zawiya database.
"""
import os, sys, re, json, glob, hashlib
from collections import Counter

BOOKS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'lirbrery_books')
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'backend', 'database')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, 'extracted_books.json')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── DDC-to-category mapping ────────────────────────────────────────
DDC_CATEGORIES = {
    '020': 'المكتبات والمعلومات',
    '050': 'الدوريات العامة',
    '110': 'الميتافيزيقا',
    '150': 'علم النفس',
    '230': 'علم الكلام',
    '240': 'الأخلاق والتصوف',
    '280': 'الفرق والمذاهب',
    '320': 'العلوم السياسية',
    '340': 'القانون',
    '413': 'المعاجم والقواميس',
    '414': 'علم اللغة',
    '415': 'النحو والصرف',
    '510': 'الرياضيات',
    '530': 'الفيزياء',
    '550': 'علوم الأرض',
    '610': 'الطب',
    '720': 'الفنون',
    '740': 'الرسم والزخرفة',
    '820': 'الشعر',
    '940': 'التاريخ',
}

def ddc_to_category(ddc_num):
    """Map DDC number to category name, or use number as fallback."""
    if ddc_num in DDC_CATEGORIES:
        return DDC_CATEGORIES[ddc_num]
    return f'تصنيف {ddc_num}'

def parse_ddc_from_title(title):
    """Extract DDC number from title like 'دواويـن الشعـر / 820' or 'علوم / 510'"""
    if not title:
        return None
    m = re.search(r'/\s*(\d{2,4})', title)
    if m:
        return m.group(1)
    return None

def content_hash(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def parse_table_rows(table):
    """Extract book data rows from a table, skipping header rows."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if not cells or not cells[0]:
            continue
        # Skip header rows (Arabic header column)
        if cells[0] in ('رقم الكتاب', 'الرقم', 'م'):
            continue
        # Row must have at least book number and title to be valid
        book_num = cells[0] if len(cells) > 0 else ''
        copies = cells[1] if len(cells) > 1 else ''
        author = cells[2] if len(cells) > 2 else ''
        title = cells[3] if len(cells) > 3 else ''
        index = cells[4] if len(cells) > 4 else ''
        
        if not title.strip():
            continue
        
        rows.append({
            'book_number': book_num.strip(),
            'copies': copies.strip(),
            'author': author.strip(),
            'title': re.sub(r'\s+', ' ', title).strip(),
            'index': index.strip(),
        })
    return rows

def process_docx(filepath):
    """Process a .docx catalog file and return (category_name, ddc_num, books)."""
    from docx import Document
    
    doc = Document(filepath)
    
    # Get category title from first paragraph
    category_title = ''
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            category_title = t
            break
    
    ddc_num = parse_ddc_from_title(category_title)
    category_name = ddc_to_category(ddc_num) if ddc_num else 'عام'
    
    if ddc_num:
        clean_title = category_title.rsplit(' /', 1)[0].strip()
    else:
        clean_title = category_title
    
    # Extract all book rows from all tables
    all_books = []
    for table in doc.tables:
        all_books.extend(parse_table_rows(table))
    
    return {
        'ddc_num': ddc_num or '',
        'category_title': clean_title,
        'category_name': category_name,
        'books': all_books,
    }, all_books

def main():
    print('=' * 60)
    print('BOOK CATALOG EXTRACTOR — Zawiya Digital Library')
    print('=' * 60)
    
    if not os.path.isdir(BOOKS_DIR):
        print(f'ERROR: Books directory not found: {BOOKS_DIR}')
        sys.exit(1)
    
    files = sorted(glob.glob(os.path.join(BOOKS_DIR, '*.docx')))
    print(f'\nFound {len(files)} .docx files in {BOOKS_DIR}')
    
    seen_hashes = set()
    all_categories = []
    all_books = []
    skipped = 0
    
    for i, fp in enumerate(files):
        fname = os.path.basename(fp)
        print(f'\n[{i+1}/{len(files)}] {fname} ({os.path.getsize(fp)} bytes)')
        
        try:
            cat_info, books = process_docx(fp)
        except Exception as e:
            print(f'  ERROR: {e}')
            skipped += 1
            continue
        
        # Deduplicate by content hash of the book list
        books_json = json.dumps(books, ensure_ascii=False)
        h = content_hash(books_json)
        if h in seen_hashes:
            print(f'  SKIP: Duplicate of previously processed file')
            skipped += 1
            continue
        seen_hashes.add(h)
        
        all_categories.append(cat_info)
        all_books.extend(books)
        
        ddc = cat_info['ddc_num']
        cat = cat_info['category_name']
        count = len(books)
        title_preview = cat_info['category_title'][:50]
        print(f'  DDC: {ddc or "?"}  |  Category: {cat}  |  Books: {count}')
        if title_preview:
            print(f'  Title: {title_preview}')
    
    # Summary
    print('\n' + '=' * 60)
    print(f'RESULTS: {len(all_categories)} categories, {len(all_books)} book entries, {skipped} skipped')
    print('=' * 60)
    
    # Build full book records for output
    book_records = []
    cat_stats = Counter()
    for ci in all_categories:
        cat_name = ci['category_name']
        ddc = ci['ddc_num']
        count = len(ci['books'])
        cat_stats[cat_name] += count
        for b in ci['books']:
            book_records.append({
                'title': b['title'],
                'author': b['author'] or 'غير معروف',
                'description': '',
                'category': cat_name,
                'ddc_number': ddc,
                'book_number': b['book_number'],
                'copies': int(b['copies']) if b['copies'].isdigit() else 1,
                'language': 'Arabic',
                'publisher': '',
                'license_type': 'Public Domain',
                'is_visible': 1,
                'is_featured': 0,
            })
    
    # Category breakdown
    print(f'\nCategory distribution:')
    for cat_name, n in cat_stats.most_common():
        ddc_nums = set()
        for ci in all_categories:
            if ci['category_name'] == cat_name and ci['ddc_num']:
                ddc_nums.add(ci['ddc_num'])
        ddc_str = ','.join(sorted(ddc_nums)) if ddc_nums else ''
        print(f'  {cat_name:25s} ({ddc_str:10s}) : {n:4d} books')
    
    # Save JSON
    output = {
        'total_books': len(book_records),
        'total_categories': len(all_categories),
        'categories': all_categories,
        'books': book_records,
    }
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f'\nJSON saved to: {OUTPUT_FILE}')
    
    # Generate seed SQL
    seed_sql(book_records, all_categories)
    
    print('\nDone!')

def seed_sql(books, categories):
    sql_file = os.path.join(OUTPUT_DIR, 'seed_extracted.sql')
    
    lines = [
        '-- Seed script for extracted library catalog books',
        '-- Generated by extract_books.py',
        f'-- {len(books)} books across {len(categories)} categories',
        '',
        "PRAGMA journal_mode=WAL;",
        '',
    ]
    
    # Gather unique category names (in order from categories list)
    seen_cats = set()
    cat_names = []
    for ci in categories:
        n = ci['category_name']
        if n not in seen_cats:
            seen_cats.add(n)
            cat_names.append(n)
    
    lines.append('-- Insert categories if they do not exist')
    lines.append('INSERT OR IGNORE INTO categories (name, slug, color) VALUES')
    cat_lines = []
    for i, name in enumerate(cat_names):
        slug = re.sub(r'[^a-zA-Z0-9\u0600-\u06ff-]', '-', name)
        slug = re.sub(r'-+', '-', slug).strip('-')
        cat_lines.append(f"    ('{name}', '{slug}', NULL)")
    lines.append(',\n'.join(cat_lines) + ';')
    
    lines.extend(['', '-- Insert books', ''])
    lines.append('INSERT OR IGNORE INTO books (')
    lines.append('    title, author, description, category, language,')
    lines.append('    publisher, license_type, is_visible, is_featured')
    lines.append(') VALUES')
    
    book_lines = []
    for b in books:
        title = b['title'].replace("'", "''")
        author = b['author'].replace("'", "''") if b['author'] else 'غير معروف'
        desc = b['description'].replace("'", "''") if b['description'] else ''
        cat = b['category'].replace("'", "''")
        lang = b['language'].replace("'", "''")
        pub = b['publisher'].replace("'", "''") if b['publisher'] else ''
        lic = b['license_type'].replace("'", "''")
        visible = b['is_visible']
        featured = b['is_featured']
        
        book_lines.append(f"    ('{title}', '{author}', '{desc}', '{cat}', '{lang}', '{pub}', '{lic}', {visible}, {featured})")
    
    lines.append(',\n'.join(book_lines) + ';')
    
    with open(sql_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f'SQL seed: {sql_file}')

if __name__ == '__main__':
    main()
